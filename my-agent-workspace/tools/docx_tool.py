"""
tools/docx_tool.py
──────────────────
Skill 4 — Write .docx files on Hao's Mac.

Install: pip install python-docx
"""

import json
import os
import subprocess
import uuid

try:
    from upsonic.tools import tool
except ImportError:
    def tool(fn): return fn

from config import WORKSPACE


def create_docx_file(
    filepath: str,
    title: str,
    sections: list[dict],
    mode: str = "create",
) -> str:
    """
    Write a .docx file on Hao's Mac using python-docx.

    mode: "create"  — start fresh (default). Overwrites if file exists.
          "append"  — open existing file and add sections at END.
          "prepend" — open existing file and insert sections at TOP.
    """
    sections_json     = json.dumps(sections, ensure_ascii=False)
    filepath_expanded = os.path.expanduser(filepath)

    script = f"""
import sys, os, json, copy
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

filepath  = {json.dumps(filepath_expanded)}
title_str = {json.dumps(title)}
sections  = json.loads({json.dumps(sections_json)})
mode      = {json.dumps(mode)}

if mode in ("append", "prepend") and os.path.exists(filepath):
    doc = Document(filepath)
else:
    doc = Document()
    p = doc.add_heading(title_str, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _add(target, secs):
    for sec in secs:
        t = sec.get("type", "paragraph")
        if t == "heading":
            target.add_heading(sec.get("text",""), level=int(sec.get("level",1)))
        elif t == "paragraph":
            target.add_paragraph(sec.get("text",""))
        elif t == "bullet":
            target.add_paragraph(sec.get("text",""), style="List Bullet")
        elif t == "table":
            h = sec.get("headers", [])
            if h:
                tbl = target.add_table(rows=1, cols=len(h))
                tbl.style = "Table Grid"
                for i, v in enumerate(h): tbl.rows[0].cells[i].text = str(v)
                for row in sec.get("rows", []):
                    cells = tbl.add_row().cells
                    for i, v in enumerate(row[:len(h)]): cells[i].text = str(v)
        elif t == "pagebreak":
            target.add_page_break()

if mode == "prepend":
    try:
        from docx import Document as _D
        tmp = _D()
        _add(tmp, sections)
        body = doc.element.body
        ref  = body.paragraphs[0]._element if body.paragraphs else None
        for elem in reversed(list(tmp.element.body)):
            if elem.tag.endswith("}}sectPr"): continue
            idx = list(body).index(ref)+1 if ref is not None else 0
            body.insert(idx, copy.deepcopy(elem))
    except Exception as e:
        print(f"prepend fallback: {{e}}")
        _add(doc, sections)
else:
    _add(doc, sections)

os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else ".", exist_ok=True)
doc.save(filepath)
print(f"OK: {{filepath}}")
"""

    script_path = os.path.join(WORKSPACE, f"_docx_temp_{uuid.uuid4().hex[:8]}.py")
    try:
        os.makedirs(os.path.dirname(script_path), exist_ok=True)
        with open(script_path, "w") as f:
            f.write(script)
        result = subprocess.run(
            ["python3", script_path],
            capture_output=True, text=True, timeout=30,
        )
        output = (result.stdout + result.stderr).strip()
        return output if output else "(docx saved — no output)"
    except subprocess.TimeoutExpired:
        return "[docx_tool] Timeout — script took over 30 seconds."
    except Exception as e:
        return f"[docx_tool] Error: {e}"
    finally:
        if os.path.exists(script_path):
            os.remove(script_path)


@tool
def docx_tool(filepath: str, title: str, sections: str, mode: str = "create") -> str:
    """
    Write a Microsoft Word (.docx) file directly on Hao's Mac.

    Args:
        filepath: Save path, e.g. "~/Desktop/report.docx"
        title:    Document title heading.
        sections: JSON string — list of section dicts, each with keys:
                  type ("heading"|"paragraph"|"bullet"|"table"|"pagebreak"),
                  text, level (for headings), headers + rows (for tables).
        mode:     "create" (default), "append", or "prepend".
    Returns:
        "OK: /path/to/file.docx" on success, or an error message.
    """
    try:
        parsed = json.loads(sections)
    except json.JSONDecodeError as e:
        return f"[docx_tool] Invalid sections JSON: {e}"
    return create_docx_file(filepath, title, parsed, mode=mode)
